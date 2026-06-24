import datetime
from io import BytesIO

import numpy as np
import pandas as pd
from engine._streamlit_shim import streamlit as st   # headless-safe; see shim docstring
import yfinance as yf
from sklearn.linear_model import LassoCV, Ridge, ElasticNetCV

# python-docx is optional — only needed by generate_docx_report(). Lazy-import
# inside that function so risk_metrics.py and downstream pages can load even
# when python-docx is not installed (2026-05-14 bootstrap recovery).


def _bootstrap_ci(
    series: pd.Series,
    stat_fn,
    n_boot: int = 1000,
    alpha: float = 0.05,
    seed: int = 42,
) -> tuple[float, float]:
    """
    Asymmetric percentile bootstrap CI.
    Resamples `series` with replacement `n_boot` times, applies `stat_fn`,
    and returns the (alpha/2, 1-alpha/2) percentiles.
    The resulting interval is naturally skewed — no normality assumption.
    """
    n = len(series)
    if n < 20:
        return (float("nan"), float("nan"))
    rng  = np.random.default_rng(seed)
    arr  = series.values
    boot = []
    for _ in range(n_boot):
        idx = rng.integers(0, n, size=n)
        val = stat_fn(pd.Series(arr[idx]))
        if np.isfinite(val):
            boot.append(val)
    if len(boot) < 50:
        return (float("nan"), float("nan"))
    return float(np.percentile(boot, alpha / 2 * 100)), float(np.percentile(boot, (1 - alpha / 2) * 100))


class StrategyAuditor:
    @staticmethod
    def run_feature_sparsity_check(X, y):
        """
        Fit regularised linear model with a strict temporal train/test split.

        Regulariser is chosen by sample size to balance bias-variance:
          n < 30   → Ridge (α=1.0): L2 prior, all-features, most stable at tiny n
          30≤n<100 → ElasticNetCV: L1+L2 mix, partial sparsity, moderate n
          n ≥ 100  → LassoCV: L1 prior, true sparsity only reliable at large n

        Rationale: Lasso's sparsity advantage requires n large enough for the
        constraint boundary to reliably separate signal from noise features.
        At n < 30 the solution is unstable — a single obs change can flip which
        features are selected (Tibshirani 1996; Fan & Li 2001).

        Split rationale: financial returns are time-ordered; random k-fold would
        introduce look-ahead bias. We use the first 80% of observations for
        training and the remaining 20% as a true out-of-sample hold-out.

        Returns
        -------
        active_features : int
        sparsity        : float   fraction of zero coefficients
        coefs           : ndarray model.coef_ on train set
        val_r2          : float | None  in-sample R² (None if n < 30)
        test_r2         : float | None  out-of-sample R² (None if n < 30)
        n_alphas        : int     number of regularisation values searched (0 for Ridge)
        """
        common_index = X.index.intersection(y.index)
        X, y = X.loc[common_index], y.loc[common_index]
        n = len(X)

        # n < 30: Ridge with structural prior — no CV, no train/test split
        if n < 30:
            model = Ridge(alpha=1.0).fit(X, y)
            coefs = model.coef_
            active_features = int(np.sum(np.abs(coefs) > 1e-10))
            sparsity = 1 - (active_features / X.shape[1])
            return active_features, sparsity, coefs, None, None, 0, "Ridge"

        # Temporal split (no shuffling — time-series integrity)
        split   = int(n * 0.8)
        X_train = X.iloc[:split];  X_test = X.iloc[split:]
        y_train = y.iloc[:split];  y_test = y.iloc[split:]

        if n < 100:
            # ElasticNetCV: L1_ratio=0.5 balances sparsity and ridge stability
            model      = ElasticNetCV(l1_ratio=0.5, cv=min(3, max(2, split // 10))).fit(X_train, y_train)
            n_alphas   = int(len(model.alphas_)) if hasattr(model, "alphas_") else 0
            model_name = "ElasticNet"
        else:
            model      = LassoCV(cv=5).fit(X_train, y_train)
            n_alphas   = int(len(model.alphas_))
            model_name = "Lasso"

        coefs           = model.coef_
        active_features = int(np.sum(np.abs(coefs) > 1e-10))

        # Post-hoc sparsity fallback: if L1 penalty zeroed out almost all features,
        # the sparse solution is degenerate — re-fit with Ridge (L2) which never
        # produces zero coefficients and is more stable on sparse financial data.
        _MIN_ACTIVE = 2
        if active_features < _MIN_ACTIVE:
            ridge_fb   = Ridge(alpha=1.0).fit(X_train, y_train)
            coefs      = ridge_fb.coef_
            active_features = int(np.sum(np.abs(coefs) > 1e-10))
            sparsity   = 1 - (active_features / X_train.shape[1])
            val_r2     = float(ridge_fb.score(X_train, y_train))
            test_r2    = float(ridge_fb.score(X_test,  y_test))
            n_alphas   = 0
            model_name = f"Ridge(fallback from {model_name})"
        else:
            sparsity = 1 - (active_features / X_train.shape[1])
            val_r2   = float(model.score(X_train, y_train))
            test_r2  = float(model.score(X_test,  y_test))

        return active_features, sparsity, coefs, val_r2, test_r2, n_alphas, model_name

    @staticmethod
    def check_optimizer_curse(val_score, test_score, n_trials):
        """
        Estimate the probability that the in-sample performance is noise.

        gap         = max(0, val_score - test_score)  — overfit magnitude
        luck_factor = log(1 + n_trials) * 0.01        — expected gap from search
        prob_noise  = 1 - exp(-gap / luck_factor)

        A large gap relative to luck_factor indicates the strategy was fitted to
        noise. When val_score and test_score are real R² values from a temporal
        split, this gives a data-driven estimate rather than a heuristic constant.
        """
        luck_factor = np.log1p(n_trials) * 0.01
        gap = max(0, val_score - test_score)
        prob_noise = 1 - np.exp(-gap / (luck_factor + 1e-6))
        return gap <= luck_factor, prob_noise


class QuantEngine:
    @staticmethod
    @st.cache_data(ttl=3600, show_spinner=False)
    def get_market_data(tickers: tuple) -> pd.DataFrame:
        """返回日收益率 DataFrame；tickers 必须是 tuple 以支持缓存。"""
        try:
            data = yf.download(list(tickers), period="1y", progress=False)["Close"]
            if isinstance(data, pd.Series):
                data = data.to_frame()
            return data.pct_change().fillna(0)
        except Exception:
            return pd.DataFrame()

    @staticmethod
    def cornish_fisher_var(returns: pd.Series, alpha: float = 0.05) -> float:
        """
        Parametric VaR adjusted for skewness and excess kurtosis via Cornish-Fisher expansion.
        More accurate than normal-distribution VaR when returns are skewed or fat-tailed.
        Returns same sign convention as np.percentile VaR (negative = loss).

        Recommended to enable (use_cf=True) once backtest sample > 200 decisions,
        at which point the skew/kurtosis estimates become stable.

        CF expansion formula:
            z_cf = z + (z²-1)·S/6 + (z³-3z)·K/24 − (2z³-5z)·S²/36
        where S = skewness, K = excess kurtosis, z = normal quantile at alpha.
        """
        mu    = float(returns.mean())
        sigma = float(returns.std())
        skew  = float(returns.skew())
        ekurt = float(returns.kurtosis())   # excess kurtosis (Fisher convention)
        z     = -1.6449                     # norm.ppf(0.05), hardcoded to avoid scipy dep
        z_cf  = (z
                 + (z**2 - 1) * skew / 6
                 + (z**3 - 3 * z) * ekurt / 24
                 - (2 * z**3 - 5 * z) * skew**2 / 36)
        return mu + z_cf * sigma

    @staticmethod
    def expected_shortfall(returns: pd.Series, alpha: float = 0.05) -> float:
        """ES (CVaR): mean of the worst α% daily returns. Basel III preferred tail risk metric.
        More informative than VaR — answers 'when things go bad, how bad on average?'"""
        cutoff = float(np.quantile(returns, alpha))
        tail = returns[returns <= cutoff]
        return float(tail.mean()) if len(tail) > 0 else cutoff

    @staticmethod
    def compute_asymmetric_risk(
        returns: pd.DataFrame, vix: float, weights: np.ndarray
    ) -> tuple[float, float, float]:
        port_ret = returns.dot(weights)
        alpha = 0.05 * (1 + (vix - 20) / 40) if vix > 20 else 0.05
        dynamic_var = float(np.quantile(port_ret, min(alpha, 0.2)))
        return dynamic_var, float(port_ret.mean() * 252), float(port_ret.std() * np.sqrt(252))

    @staticmethod
    def get_realtime_vix() -> float:
        # Try fast_info first — single scalar, no MultiIndex ambiguity
        try:
            fi = yf.Ticker("^VIX").fast_info
            val = float(fi.last_price or fi.regular_market_previous_close or 0)
            if 5.0 < val < 120.0:   # VIX has never exceeded ~90 in history
                return round(val, 2)
        except Exception:
            pass
        # Fallback: daily close from download
        try:
            vix_data = yf.download("^VIX", period="5d", progress=False,
                                   auto_adjust=True, multi_level_index=False)
            if not vix_data.empty:
                val = float(vix_data["Close"].dropna().iloc[-1])
                if 5.0 < val < 120.0:
                    return round(val, 2)
        except Exception:
            pass
        return 20.0


def compute_quant_metrics(tickers: tuple, vix: float) -> dict:
    """
    Compute quantitative metrics for a basket of tickers.

    This is the canonical quant-metrics computation shared by:
      - Quant Audit agent (engine/agent.py research_node)
      - Sector Analysis pipeline (ui/tabs.py _run_sector_analysis)

    Returns a metrics dict on success, empty dict if market data is unavailable.

    Keys returned:
      d_var       float      dynamic VaR (daily, negative = loss)
      a_ret       float      annualised return (equal-weight portfolio)
      a_vol       float      annualised volatility
      sharpe      float      annualised Sharpe ratio (a_ret / a_vol)
      sharpe_ci   (lo, hi)   95% asymmetric bootstrap CI for Sharpe
      mom_1m      float|None 1-month cumulative return (21 trading days)
      mom_3m      float|None 3-month cumulative return (63 trading days)
      mom_6m      float|None 6-month cumulative return (126 trading days)
      skewness    float      sample skewness of daily portfolio returns
      excess_kurt float      sample excess kurtosis (>0 = fat tails)
      var_cf      float      Cornish-Fisher VaR (skew/kurt adjusted)
      var_ci      (lo, hi)   95% asymmetric bootstrap CI for historical VaR
      p_noise     float      optimizer-curse noise probability (0-1); 1.0 if sample < 60
      val_r2      float|None in-sample R² from temporal-split Lasso
      test_r2     float|None out-of-sample R² from temporal-split Lasso
      n_alphas    int        number of regularisation candidates searched
      active      int        number of non-zero Lasso coefficients
      sparsity    float      fraction of zero coefficients
      coefs       ndarray    Lasso coefficients (same fit as val_r2/test_r2)
      X           DataFrame  feature matrix used for Lasso (for downstream display)

    coefs and X are included so callers (e.g. agent.py research_node) do not need
    to re-run LassoCV independently — two independent fits would produce inconsistent
    coefficients due to non-deterministic CV fold selection.

    Distribution stats (sharpe_ci, skewness, excess_kurt, var_cf, var_ci) are computed
    from the same port_ret series and replace the need to call AnalyticsEngine.calculate_metrics()
    separately in Tab4 — both tabs now share a single canonical computation.
    """
    returns = QuantEngine.get_market_data(tickers)
    if returns.empty:
        return {}

    n_assets = len(returns.columns)
    weights  = np.full(n_assets, 1.0 / n_assets)

    d_var, a_ret, a_vol = QuantEngine.compute_asymmetric_risk(returns, vix, weights)

    port_ret = returns.dot(weights)
    cum      = (1 + port_ret).cumprod()
    n        = len(cum)

    def _mom(window: int):
        return float(cum.iloc[-1] / cum.iloc[-1 - window] - 1) if n > window else None

    # ── Distribution statistics ───────────────────────────────────────────────
    # Computed from port_ret so Tab3 and Tab4 share identical inputs.
    sharpe     = float(a_ret / a_vol) if a_vol else 0.0
    skewness   = float(port_ret.skew())
    excess_kurt = float(port_ret.kurtosis())
    var_cf     = QuantEngine.cornish_fisher_var(port_ret, alpha=0.05)

    _alpha_ci = 0.05
    def _var_fn(s: pd.Series) -> float:
        return float(np.percentile(s, _alpha_ci * 100))
    def _sharpe_fn(s: pd.Series) -> float:
        mu_a = s.mean() * 252
        sd_a = s.std() * np.sqrt(252)
        return float(mu_a / sd_a) if sd_a > 1e-9 else 0.0

    var_ci    = _bootstrap_ci(port_ret, _var_fn)
    sharpe_ci = _bootstrap_ci(port_ret, _sharpe_fn)
    es_5pct   = QuantEngine.expected_shortfall(port_ret, alpha=0.05)

    X = pd.concat([
        returns.shift(1).add_suffix("_lag1"),
        pd.DataFrame({
            "mom_21d": (1 + port_ret).rolling(21, min_periods=5).apply(np.prod) - 1,
            "mom_63d": (1 + port_ret).rolling(63, min_periods=10).apply(np.prod) - 1,
        }),
    ], axis=1).fillna(0)

    y = returns.iloc[:, 0]
    active, sparsity, coefs, val_r2, test_r2, n_alphas, model_name = (
        StrategyAuditor.run_feature_sparsity_check(X, y)
    )

    if val_r2 is not None:
        _, p_noise = StrategyAuditor.check_optimizer_curse(val_r2, test_r2, n_alphas)
    else:
        p_noise = 1.0   # sample too small → maximum uncertainty

    return {
        "d_var":        d_var,
        "a_ret":        a_ret,
        "a_vol":        a_vol,
        "sharpe":       sharpe,
        "sharpe_ci":    sharpe_ci,
        "mom_1m":       _mom(21),
        "mom_3m":       _mom(63),
        "mom_6m":       _mom(126),
        "skewness":     skewness,
        "excess_kurt":  excess_kurt,
        "var_cf":       var_cf,
        "es_5pct":      es_5pct,
        "var_ci":       var_ci,
        "p_noise":      p_noise,
        "val_r2":       val_r2,
        "test_r2":      test_r2,
        "n_alphas":     n_alphas,
        "active":       active,
        "sparsity":     sparsity,
        "coefs":        coefs,
        "model_name":   model_name,
        "X":            X,
    }


class AnalyticsEngine:
    @staticmethod
    def calculate_metrics(
        df: pd.DataFrame,
        confidence_level: float = 0.95,
        use_cf: bool = False,
        n_boot: int = 1000,
    ) -> dict | None:
        """
        Compute portfolio risk metrics with asymmetric bootstrap confidence intervals.

        Parameters
        ----------
        df            : price DataFrame (Close prices, one column per ticker)
        confidence_level : CI coverage, default 0.95 (95%)
        use_cf        : if True, replace historical VaR with Cornish-Fisher adjusted VaR.
                        Recommended only when sample > 200 obs (skew/kurt estimates stabilise).
                        Default False — will be enabled progressively as backtest data grows.
        n_boot        : bootstrap resamples, default 1000

        Returns additional keys vs. legacy version:
            var_ci      : (lo, hi) bootstrap CI for VaR  — asymmetric (left-tail wider)
            sharpe_ci   : (lo, hi) bootstrap CI for Sharpe — asymmetric (right-skewed when SR>0)
            var_cf      : Cornish-Fisher VaR (always computed for reference, even if use_cf=False)
            skewness    : sample skewness of daily returns
            excess_kurt : sample excess kurtosis
        """
        if df is None or df.empty:
            return None
        returns = df.pct_change().dropna()
        portfolio_returns = returns.mean(axis=1) if len(returns.columns) > 1 else returns.iloc[:, 0]
        if portfolio_returns.empty:
            return None

        alpha    = 1 - confidence_level                          # 0.05 for 95% CI
        var_hist = float(np.percentile(portfolio_returns, alpha * 100))
        var_cf   = QuantEngine.cornish_fisher_var(portfolio_returns, alpha=alpha)
        var_val  = var_cf if use_cf else var_hist

        mean_ret = portfolio_returns.mean() * 252
        std_ret  = portfolio_returns.std() * np.sqrt(252)
        sharpe   = float(mean_ret / std_ret) if std_ret != 0 else 0.0

        # ── Bootstrap CIs (asymmetric by construction) ────────────────────────
        def _var_fn(s: pd.Series) -> float:
            return float(np.percentile(s, alpha * 100))

        def _sharpe_fn(s: pd.Series) -> float:
            mu_a = s.mean() * 252
            sd_a = s.std() * np.sqrt(252)
            return float(mu_a / sd_a) if sd_a > 1e-9 else 0.0

        var_ci    = _bootstrap_ci(portfolio_returns, _var_fn,    n_boot=n_boot, alpha=alpha)
        sharpe_ci = _bootstrap_ci(portfolio_returns, _sharpe_fn, n_boot=n_boot, alpha=alpha)

        return {
            "var":          var_val,
            "var_hist":     var_hist,
            "var_cf":       var_cf,
            "es_5pct":      QuantEngine.expected_shortfall(portfolio_returns, alpha=alpha),
            "var_ci":       var_ci,          # (lo, hi) — left tail naturally wider
            "sharpe":       sharpe,
            "sharpe_ci":    sharpe_ci,       # (lo, hi) — right-skewed when SR > 0
            "volatility":   float(std_ret),
            "skewness":     float(portfolio_returns.skew()),
            "excess_kurt":  float(portfolio_returns.kurtosis()),
            "returns_series": portfolio_returns,
        }


def build_quant_context(sector_etf: str, as_of_date) -> str:
    """
    Build a structured [QUANT_CONTEXT] block for injection into AI prompts.
    All indicators use data strictly available BEFORE as_of_date (no look-ahead).

    Includes: price momentum (1M/3M/6M), relative strength vs SPY,
              RSI(14), Bollinger Band position, credit spread proxy (HYG-IEF).

    Returns empty string on failure so callers can safely skip injection.
    """
    import logging as _logging
    _log = _logging.getLogger(__name__)
    try:
        cutoff = pd.Timestamp(as_of_date) - pd.Timedelta(days=1)
        start  = (cutoff - pd.Timedelta(days=420)).strftime("%Y-%m-%d")
        end    = cutoff.strftime("%Y-%m-%d")

        raw = yf.download(
            [sector_etf, "SPY", "HYG", "IEF"],
            start=start, end=end, progress=False, auto_adjust=True,
        )
        prices = raw["Close"] if isinstance(raw.columns, pd.MultiIndex) else raw
        if isinstance(prices, pd.Series):
            prices = prices.to_frame(sector_etf)

        etf = prices.get(sector_etf, pd.Series(dtype=float)).dropna()
        spy = prices.get("SPY",      pd.Series(dtype=float)).dropna()
        hyg = prices.get("HYG",      pd.Series(dtype=float)).dropna()
        ief = prices.get("IEF",      pd.Series(dtype=float)).dropna()

        if len(etf) < 22:
            return ""

        lines = ["[QUANT_CONTEXT]", f"板块ETF: {sector_etf}"]

        # ── Momentum ─────────────────────────────────────────────────────────
        def _mom(n: int) -> float | None:
            return float(etf.iloc[-1] / etf.iloc[-n] - 1) if len(etf) > n else None

        mom_parts = []
        for label, n in [("1M", 22), ("3M", 63), ("6M", 126)]:
            v = _mom(n)
            if v is not None:
                mom_parts.append(f"{label} {v:+.1%}")
        if mom_parts:
            lines.append(f"价格动量：{' | '.join(mom_parts)}")

        # ── Relative strength vs SPY (1M) ─────────────────────────────────
        common = etf.index.intersection(spy.index)
        if len(common) > 22:
            e, s = etf.loc[common], spy.loc[common]
            rel = float(e.iloc[-1]/e.iloc[-22] - 1) - float(s.iloc[-1]/s.iloc[-22] - 1)
            lines.append(
                f"相对强弱 vs SPY（1M）：{rel:+.1%}（{'跑赢' if rel > 0 else '跑输'}大盘）"
            )

        # ── RSI(14) ───────────────────────────────────────────────────────
        delta = etf.diff()
        gain  = delta.clip(lower=0).rolling(14).mean()
        loss  = (-delta.clip(upper=0)).rolling(14).mean()
        rsi   = float(100 - 100 / (1 + gain.iloc[-1] / max(loss.iloc[-1], 1e-9)))
        rsi_note = "超买区间·需警惕回调" if rsi > 70 else ("超卖区间·可能反弹" if rsi < 30 else "中性区间")
        lines.append(f"RSI(14)：{rsi:.0f} — {rsi_note}")

        # ── Bollinger Band position (20-day) ──────────────────────────────
        bb_mid = etf.rolling(20).mean()
        bb_std = etf.rolling(20).std()
        bb_rng = (bb_mid + 2*bb_std).iloc[-1] - (bb_mid - 2*bb_std).iloc[-1]
        if bb_rng > 0:
            bb_pos = float((etf.iloc[-1] - (bb_mid - 2*bb_std).iloc[-1]) / bb_rng)
            bb_pos = max(0.0, min(1.0, bb_pos))
            bb_note = (
                "接近上轨·强势但需警惕" if bb_pos > 0.80
                else ("接近下轨·可能超卖" if bb_pos < 0.20 else "中性区间")
            )
            lines.append(f"布林带位置：{bb_pos:.0%} — {bb_note}")

        # ── Credit spread proxy (HYG - IEF 20-day return spread) ─────────
        ci = hyg.index.intersection(ief.index)
        if len(ci) > 21:
            cs = float(hyg.loc[ci].pct_change(20).iloc[-1] -
                       ief.loc[ci].pct_change(20).iloc[-1])
            cs_note = "收窄→风险偏好改善" if cs > 0 else "扩大→信用风险上升"
            lines.append(f"信用利差(HYG-IEF 20日价差)：{cs:+.2%} — {cs_note}")

        lines.append("[/QUANT_CONTEXT]")
        return "\n".join(lines)

    except Exception as exc:
        import logging as _lg
        _lg.getLogger(__name__).warning("build_quant_context failed for %s: %s", sector_etf, exc)
        return ""


def compute_state_vector(sector_etf: str, as_of_date) -> dict[str, str]:
    """
    Compute a discrete regime state vector for state-change detection in backtest sampling.
    Compares the vector between consecutive dates; only triggers a new analysis when
    at least one dimension changes (Layer 2 deduplication).

    Dimensions:
        momentum_regime : "上升" | "震荡" | "下降"  (1M return threshold ±2%)
        rsi_zone        : "超买" | "中性" | "超卖"  (RSI 70/30)
        vol_regime      : "高波动" | "正常" | "低波动" (20d vol vs 1y historical)

    Returns a dict of string labels. On failure returns all "unknown" (treated as changed).
    """
    state: dict[str, str] = {
        "momentum_regime": "unknown",
        "rsi_zone":        "unknown",
        "vol_regime":      "unknown",
    }
    try:
        cutoff = pd.Timestamp(as_of_date) - pd.Timedelta(days=1)
        start  = (cutoff - pd.Timedelta(days=260)).strftime("%Y-%m-%d")
        end    = cutoff.strftime("%Y-%m-%d")

        raw  = yf.download(sector_etf, start=start, end=end, progress=False, auto_adjust=True)
        etf  = raw["Close"] if isinstance(raw.columns, pd.MultiIndex) else raw["Close"]
        if isinstance(etf, pd.DataFrame):
            etf = etf.iloc[:, 0]
        etf = etf.dropna()

        if len(etf) < 30:
            return state

        # Momentum regime
        mom = float(etf.iloc[-1] / etf.iloc[-22] - 1) if len(etf) > 22 else 0.0
        state["momentum_regime"] = "上升" if mom > 0.02 else ("下降" if mom < -0.02 else "震荡")

        # RSI zone
        d    = etf.diff()
        gain = d.clip(lower=0).rolling(14).mean()
        loss = (-d.clip(upper=0)).rolling(14).mean()
        rsi  = float(100 - 100 / (1 + gain.iloc[-1] / max(loss.iloc[-1], 1e-9)))
        state["rsi_zone"] = "超买" if rsi > 70 else ("超卖" if rsi < 30 else "中性")

        # Volatility regime (20d vs full-sample historical)
        rets    = etf.pct_change().dropna()
        vol_20d = float(rets.tail(20).std() * (252 ** 0.5))
        vol_1y  = float(rets.std() * (252 ** 0.5))
        if vol_1y > 0:
            state["vol_regime"] = (
                "高波动" if vol_20d > vol_1y * 1.3
                else ("低波动" if vol_20d < vol_1y * 0.7 else "正常")
            )

    except Exception as exc:
        import logging as _lg
        _lg.getLogger(__name__).warning("compute_state_vector failed for %s: %s", sector_etf, exc)

    return state


@st.cache_data(ttl=3600, show_spinner=False)
def get_valuation_snapshot(ticker: str) -> str:
    """
    Fetch key valuation & positioning metrics via yfinance Ticker.info.
    Returns a formatted string for injection into analysis prompts.
    Falls back gracefully if data is unavailable.
    """
    try:
        info = yf.Ticker(ticker).info
        if not info or info.get("quoteType") == "NONE":
            return "估值数据暂不可用"

        lines: list[str] = []

        # ── Price positioning ─────────────────────────────────────────────────
        price  = info.get("currentPrice") or info.get("regularMarketPrice")
        high52 = info.get("fiftyTwoWeekHigh")
        low52  = info.get("fiftyTwoWeekLow")
        ma200  = info.get("twoHundredDayAverage")
        ma50   = info.get("fiftyDayAverage")

        if price and high52 and low52 and high52 > low52:
            pos = (price - low52) / (high52 - low52) * 100
            if pos > 70:
                pos_note = "接近52周高位 → 利好可能已较充分定价，需警惕回调风险"
            elif pos > 30:
                pos_note = "处于52周中性区间 → 定价适中，方向性信号尚不明确"
            else:
                pos_note = "接近52周低位 → 市场预期已偏悲观，利空可能已充分定价"
            lines.append(f"52周价格分位: {pos:.0f}%  {pos_note}")

        if price and ma200:
            gap = (price - ma200) / ma200 * 100
            if gap > 10:
                ma_note = "大幅高于200日均线，趋势强势但估值溢价需关注"
            elif gap > 0:
                ma_note = "位于200日均线上方，趋势偏多"
            elif gap > -10:
                ma_note = "跌破200日均线，趋势转弱"
            else:
                ma_note = "大幅低于200日均线，趋势弱势"
            lines.append(f"相对200日均线: {gap:+.1f}%  ({ma_note})")

        if price and ma50 and ma200:
            lines.append(
                f"均线结构: 50日={ma50:.2f}  200日={ma200:.2f}  "
                f"{'金叉(50>200,多头排列)' if ma50 > ma200 else '死叉(50<200,空头排列)'}"
            )

        # ── Valuation multiples ───────────────────────────────────────────────
        fpe  = info.get("forwardPE")
        tpe  = info.get("trailingPE")
        pb   = info.get("priceToBook")
        beta = info.get("beta")
        ps   = info.get("priceToSalesTrailingTwelveMonths")

        if fpe  is not None: lines.append(f"前瞻P/E: {fpe:.1f}x  (衡量市场对未来盈利增长的定价)")
        if tpe  is not None: lines.append(f"滚动P/E: {tpe:.1f}x")
        if pb   is not None: lines.append(f"市账率P/B: {pb:.2f}x")
        if ps   is not None: lines.append(f"市销率P/S: {ps:.2f}x")
        if beta is not None: lines.append(f"Beta(vs SPY): {beta:.2f}  (系统性风险敞口)")

        # ── Growth signals ────────────────────────────────────────────────────
        eg = info.get("earningsGrowth")
        rg = info.get("revenueGrowth")
        if eg is not None:
            lines.append(
                f"盈利增速(YoY): {eg:+.1%}  "
                f"({'盈利加速扩张' if eg > 0.15 else ('温和增长' if eg > 0 else '盈利收缩，需关注下修风险')})"
            )
        if rg is not None:
            lines.append(f"营收增速(YoY): {rg:+.1%}")

        return "\n".join(lines) if lines else "估值数据暂不可用"

    except Exception:
        return "估值数据暂不可用"


def fetch_raw_data(tickers: list, period: str = "1y") -> pd.DataFrame:
    """抓取历史收盘价，失败时返回空 DataFrame。"""
    if not tickers:
        return pd.DataFrame()
    try:
        data = yf.download(tickers, period=period, interval="1d", progress=False)
        close = data["Close"] if isinstance(data.columns, pd.MultiIndex) else data[["Close"]]
        return close.ffill().dropna()
    except Exception:
        return pd.DataFrame()


def generate_docx_report(content: str, title: str = "Investment Memo") -> BytesIO:
    from docx import Document   # lazy: only needed when this function is called
    doc = Document()
    doc.add_heading("Macro Alpha Intelligence", 0)
    doc.add_heading(title, level=1)
    doc.add_paragraph(f"Report Date: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}")
    for line in content.split("\n"):
        doc.add_paragraph(line)
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def _register_cjk_fonts() -> tuple[str, str]:
    """
    Register CJK-capable fonts and return (regular_name, bold_name).
    Falls back gracefully on non-Windows / cloud environments.
    """
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    candidates = [
        ("C:/Windows/Fonts/msyh.ttc",   "C:/Windows/Fonts/msyhbd.ttc",   0),
        ("C:/Windows/Fonts/simhei.ttf",  "C:/Windows/Fonts/simhei.ttf",   None),
        ("C:/Windows/Fonts/simsun.ttc",  "C:/Windows/Fonts/simsun.ttc",   0),
    ]
    for reg_path, bold_path, idx in candidates:
        try:
            kwargs_r = {"subfontIndex": idx} if idx is not None else {}
            kwargs_b = {"subfontIndex": idx} if idx is not None else {}
            pdfmetrics.registerFont(TTFont("CJKReg",  reg_path,  **kwargs_r))
            pdfmetrics.registerFont(TTFont("CJKBold", bold_path, **kwargs_b))
            return "CJKReg", "CJKBold"
        except Exception:
            continue
    return "Helvetica", "Helvetica-Bold"


def generate_pdf_report(
    content: str,
    title: str = "Investment Research Memo",
    metadata: dict | None = None,
) -> BytesIO:
    """
    Generate a professional investment-research-grade PDF with CJK font support.
    Parses markdown headings (###/##), bullets, and numbered lists.
    """
    from reportlab.lib.colors import HexColor, white
    from reportlab.lib.enums import TA_LEFT, TA_JUSTIFY
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import HRFlowable, Paragraph, SimpleDocTemplate, Spacer

    FONT_REG, FONT_BOLD = _register_cjk_fonts()

    C_PRIMARY = HexColor("#0B1F3A")
    C_ACCENT  = HexColor("#1A56DB")
    C_TEXT    = HexColor("#1E293B")
    C_MUTED   = HexColor("#64748B")
    C_BORDER  = HexColor("#E2E8F0")

    S_TITLE = ParagraphStyle(
        "title", fontName=FONT_BOLD, fontSize=17,
        textColor=C_PRIMARY, spaceBefore=4, spaceAfter=6, leading=24,
    )
    S_META = ParagraphStyle(
        "meta", fontName=FONT_REG, fontSize=8,
        textColor=C_MUTED, spaceAfter=10, leading=13,
    )
    S_SECTION = ParagraphStyle(
        "section", fontName=FONT_BOLD, fontSize=11,
        textColor=C_PRIMARY, spaceBefore=16, spaceAfter=4, leading=16,
    )
    S_BODY = ParagraphStyle(
        "body", fontName=FONT_REG, fontSize=9.5,
        textColor=C_TEXT, spaceAfter=5, leading=15, alignment=TA_JUSTIFY,
    )
    S_BULLET = ParagraphStyle(
        "bullet", fontName=FONT_REG, fontSize=9.5,
        textColor=C_TEXT, spaceAfter=3, leading=14, leftIndent=16,
    )

    now = datetime.datetime.now()

    def _header_footer(cv, doc):
        cv.saveState()
        w, h = A4
        cv.setFillColor(C_PRIMARY)
        cv.rect(0, h - 2.0 * cm, w, 2.0 * cm, fill=1, stroke=0)
        cv.setFillColor(C_ACCENT)
        cv.rect(0, h - 2.05 * cm, w, 0.07 * cm, fill=1, stroke=0)
        cv.setFillColor(white)
        cv.setFont(FONT_BOLD, 12)
        cv.drawString(1.5 * cm, h - 1.1 * cm, "MACRO ALPHA PRO")
        cv.setFont(FONT_REG, 7)
        cv.drawString(1.5 * cm, h - 1.55 * cm, "INVESTMENT RESEARCH  ·  NUS MSBA")
        cv.setFont(FONT_REG, 7)
        cv.drawRightString(w - 1.5 * cm, h - 1.1 * cm, now.strftime("%Y-%m-%d  %H:%M"))
        cv.drawRightString(w - 1.5 * cm, h - 1.55 * cm, "CONFIDENTIAL")
        cv.setStrokeColor(C_BORDER)
        cv.setLineWidth(0.4)
        cv.line(1.5 * cm, 1.15 * cm, w - 1.5 * cm, 1.15 * cm)
        cv.setFillColor(C_MUTED)
        cv.setFont(FONT_REG, 7)
        cv.drawString(1.5 * cm, 0.75 * cm,
                      "For Educational Purposes Only  ·  Not Investment Advice")
        cv.drawRightString(w - 1.5 * cm, 0.75 * cm, f"Page {doc.page}")
        cv.restoreState()

    buf = BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=1.5 * cm, rightMargin=1.5 * cm,
        topMargin=2.6 * cm, bottomMargin=1.8 * cm,
        title=title,
    )

    story: list = []

    # ── Title block ───────────────────────────────────────────────────────────
    story.append(Paragraph(title.upper(), S_TITLE))

    meta_parts = [f"Report Date: {now.strftime('%B %d, %Y  %H:%M')}"]
    if metadata:
        meta_parts += [f"{k}: {v}" for k, v in metadata.items()]
    story.append(Paragraph("  ·  ".join(meta_parts), S_META))
    story.append(HRFlowable(width="100%", thickness=0.8, color=C_ACCENT, spaceAfter=12))

    # ── Body: parse markdown-ish text ─────────────────────────────────────────
    import re as _re
    _html_tag = _re.compile(r"<[^>]+>")
    # Covers emoticons, misc symbols, dingbats, enclosed chars, transport symbols,
    # supplemental symbols — everything that causes garbled output in PDF fonts.
    _emoji = _re.compile(
        "["
        "\U0001F000-\U0001FFFF"
        "\U00002600-\U000027BF"
        "\U00002B00-\U00002BFF"
        "\U0000FE00-\U0000FE0F"
        "\U0001FA00-\U0001FA9F"
        "\U00003200-\U000032FF"
        "]+",
        flags=_re.UNICODE,
    )

    def _clean(text: str) -> str:
        """Strip HTML tags, markdown bold/italic markers, and emoji."""
        text = _html_tag.sub("", text)
        text = _emoji.sub("", text)
        return text.replace("**", "").replace("__", "").replace("*", "").strip()

    for raw in content.split("\n"):
        line = raw.strip()
        if not line:
            story.append(Spacer(1, 5))
            continue
        # Skip lines that are pure HTML blocks (e.g. error message fragments)
        if line.startswith("<") and line.endswith(">"):
            continue

        if line.startswith("### ") or line.startswith("## "):
            story.append(Paragraph(_clean(line.lstrip("#")), S_SECTION))
            story.append(HRFlowable(
                width="100%", thickness=0.4, color=C_BORDER, spaceAfter=4,
            ))
        elif line.startswith("- ") or line.startswith("• "):
            story.append(Paragraph(f"•  {_clean(line[2:])}", S_BULLET))
        elif len(line) > 2 and line[0].isdigit() and line[1] in ".、":
            story.append(Paragraph(f"{line[0]}.  {_clean(line[2:])}", S_BULLET))
        else:
            cleaned = _clean(line)
            if cleaned:
                story.append(Paragraph(cleaned, S_BODY))

    doc.build(story, onFirstPage=_header_footer, onLaterPages=_header_footer)
    buf.seek(0)
    return buf
